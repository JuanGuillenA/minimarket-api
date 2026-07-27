from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\ASUS\Documents\GitHub\minimarket-api")
ASSETS = Path(r"C:\Users\ASUS\AppData\Local\Temp\minimarket-report-shots")
OUTPUT = ROOT / "reports" / "TFINAL_GrupoXX_Apellido1Nombre1.docx"

# Simple academic preset: Arial, black hierarchy, white pages and restrained
# gray table lines. The original content is preserved without decorative chrome.
INK = "202124"
GREEN = "202124"
EMERALD = "4A4A4A"
MINT = "FFFFFF"
AMBER = "5F6368"
AMBER_LIGHT = "FFFFFF"
GRAY = "5F6368"
LIGHT_GRAY = "F5F5F5"
BORDER = "DADCE0"
WHITE = "FFFFFF"
RED = "5F6368"


def set_run_font(
    run,
    *,
    name: str = "Arial",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=4):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "bottom", "insideH"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
    for edge in ("left", "right", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_run_font(run, size=9, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False
    doc.settings.odd_and_even_pages_header_footer = False

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, before, after, color in (
        ("Heading 1", 20, 20, 6, INK),
        ("Heading 2", 16, 18, 6, INK),
        ("Heading 3", 14, 16, 4, GRAY),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = style_name != "Heading 3"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Arial"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(GRAY)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_kicker(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    set_run_font(run, size=9, color=GRAY, bold=True)
    return p


def add_title(doc: Document, text: str, size=28, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=size, color=INK, bold=True)
    return p


def add_body(doc: Document, text: str, *, bold_label: str | None = None, color=INK):
    p = doc.add_paragraph()
    if bold_label:
        r1 = p.add_run(bold_label)
        set_run_font(r1, size=11, color=color, bold=True)
        r2 = p.add_run(text)
        set_run_font(r2, size=11, color=color)
    else:
        set_run_font(p.add_run(text), size=11, color=color)
    return p


def add_callout(doc: Document, title: str, body: str, fill=MINT, accent=EMERALD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run(title), size=11, color=INK, bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.15)
    p2.paragraph_format.space_after = Pt(8)
    set_run_font(p2.add_run(body), size=11, color=INK)
    return p2


def add_fixed_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
    *,
    font_size=9.5,
):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(header), size=font_size, color=GREEN, bold=True)
    for row in rows:
        cells = table.add_row().cells
        set_table_geometry(table, widths_dxa)
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(value), size=font_size, color=INK)
    return table


def add_screenshot(doc: Document, filename: str, caption: str, figure_no: int):
    path = ASSETS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(6.25))
    caption_p = doc.add_paragraph(
        f"Figura {figure_no}. {caption}", style="Caption"
    )
    caption_p.paragraph_format.keep_with_next = True


def add_module_page(
    doc: Document,
    *,
    number: str,
    kicker: str,
    title: str,
    lead: str,
    image: str,
    caption: str,
    purpose: str,
    actions: str,
    data: str,
    figure_no: int,
    note: str | None = None,
):
    doc.add_page_break()
    add_kicker(doc, kicker)
    add_title(doc, f"{number}  {title}", size=24)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run(lead), size=11.5, color=GRAY)
    add_screenshot(doc, image, caption, figure_no)
    add_fixed_table(
        doc,
        ["Aspecto", "Descripción"],
        [
            ["Propósito", purpose],
            ["Acciones", actions],
            ["Información", data],
        ],
        [2400, 6960],
        font_size=9.4,
    )
    if note:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run("Nota: "), size=9.5, color=EMERALD, bold=True)
        set_run_font(p.add_run(note), size=9.5, color=GRAY)


def add_cover(doc: Document):
    university = doc.add_paragraph()
    university.alignment = WD_ALIGN_PARAGRAPH.CENTER
    university.paragraph_format.space_before = Pt(10)
    university.paragraph_format.space_after = Pt(2)
    set_run_font(
        university.add_run("UNIVERSIDAD POLITÉCNICA SALESIANA"),
        size=15,
        color=INK,
        bold=True,
    )
    country = doc.add_paragraph()
    country.alignment = WD_ALIGN_PARAGRAPH.CENTER
    country.paragraph_format.space_after = Pt(34)
    set_run_font(country.add_run("Ecuador"), size=10, color=GRAY)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(8)
    set_run_font(
        kicker.add_run("PRÁCTICA FINAL - DESARROLLO WEB"),
        size=10,
        color=GRAY,
        bold=True,
    )
    add_title(
        doc,
        "Frontend para la gestión de un minimarket",
        size=26,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    set_run_font(
        subtitle.add_run("Aplicación React integrada con una API REST de operación comercial"),
        size=12,
        color=GRAY,
    )

    add_screenshot(
        doc,
        "01-dashboard.png",
        "Vista general del frontend Mercado Uno.",
        1,
    )

    meta = doc.add_table(rows=4, cols=2)
    set_table_geometry(meta, [2500, 6860])
    set_table_borders(meta)
    rows = [
        ("Grupo", "XX"),
        ("Integrantes", "[COMPLETAR NOMBRES Y APELLIDOS]"),
        ("Asignatura / docente", "[COMPLETAR]"),
        ("Fecha de entrega", "[COMPLETAR]"),
    ]
    for idx, (label, value) in enumerate(rows):
        left, right = meta.rows[idx].cells
        pl = left.paragraphs[0]
        pl.paragraph_format.space_after = Pt(0)
        set_run_font(pl.add_run(label), size=9.5, color=INK, bold=True)
        pr = right.paragraphs[0]
        pr.paragraph_format.space_after = Pt(0)
        value_color = RED if "[COMPLETAR" in value else INK
        set_run_font(pr.add_run(value), size=9.5, color=value_color)


def add_contents_page(doc: Document):
    doc.add_page_break()
    add_kicker(doc, "Documento de entrega")
    add_title(doc, "Qué debe presentar el equipo", size=25)
    add_callout(
        doc,
        "Sí, el informe es necesario",
        "La guía solicita un PDF con la evidencia de la práctica, además del enlace al repositorio. "
        "Este documento está preparado para convertirse en ese archivo final después de completar "
        "los datos del grupo y revisar que el enlace del repositorio sea público o accesible.",
        fill=AMBER_LIGHT,
        accent="8A5C00",
    )
    doc.add_heading("Entregables identificados en la guía", level=2)
    add_fixed_table(
        doc,
        ["Entregable", "Qué incluir", "Estado"],
        [
            ["PDF individual", "Informe técnico con capturas, explicación, pruebas y conclusiones.", "Preparado"],
            ["Repositorio", "Código del backend y frontend, README y evidencia del trabajo colaborativo.", "En el proyecto"],
            ["Exposición", "Recorrido ordenado de problema, solución, demostración y resultados.", "Guion incluido"],
            ["Aplicación", "Frontend funcional que consume los servicios desarrollados en la práctica previa.", "Implementado"],
        ],
        [2200, 5260, 1900],
    )
    doc.add_heading("Contenido", level=2)
    add_fixed_table(
        doc,
        ["Sección", "Página"],
        [
            ["1. Resumen, objetivos y alcance", "3"],
            ["2. Arquitectura de la solución", "3"],
            ["3-10. Recorrido por los módulos", "4-11"],
            ["11. Integración y pruebas", "12"],
            ["12. Conclusiones y guion de exposición", "13"],
        ],
        [7800, 1560],
    )
    doc.add_heading("Antes de entregar", level=2)
    checklist = [
        ("Completar datos. ", "Reemplazar los campos pendientes de la portada."),
        ("Agregar repositorio. ", "Agregar el enlace definitivo del repositorio en la última página."),
        ("Revisar commits. ", "Verificar que cada integrante tenga commits visibles; la rúbrica solicita al menos 20 en total."),
        ("Entregar. ", "Exportar el archivo con el nombre indicado por la guía y subirlo de forma individual."),
    ]
    for label, text in checklist:
        item = add_body(doc, text, bold_label=label)
        item.paragraph_format.space_after = Pt(3)
        item.paragraph_format.line_spacing = 1.0


def add_summary_architecture_page(doc: Document):
    doc.add_page_break()
    add_kicker(doc, "Visión del proyecto")
    add_title(doc, "1  Resumen, objetivos y alcance", size=24)
    add_body(
        doc,
        "Se desarrolló una interfaz web para administrar la operación cotidiana de un minimarket. "
        "La aplicación centraliza ventas, catálogo, existencias, abastecimiento, clientes, equipo y reportes "
        "en una experiencia coherente y adaptable a escritorio, tableta y móvil."
    )
    add_fixed_table(
        doc,
        ["Objetivo", "Resultado"],
        [
            ["Principal", "Consumir desde React los servicios REST desarrollados en el backend de la práctica anterior."],
            ["Operativo", "Permitir que una persona encuentre y ejecute las tareas principales con pocos pasos."],
            ["Técnico", "Mantener el frontend separado del backend, con una capa de cliente API reutilizable."],
            ["Continuidad", "Mostrar datos de demostración si la API no está disponible, sin ocultar ese estado."],
        ],
        [2500, 6860],
    )
    add_callout(
        doc,
        "Alcance de seguridad",
        "La pantalla de usuarios y roles forma parte de la interfaz, pero la autenticación y el endurecimiento "
        "de seguridad no se implementaron en esta iteración por decisión de alcance. Deben considerarse trabajo futuro.",
        fill=LIGHT_GRAY,
        accent=GREEN,
    )

    doc.add_heading("2  Arquitectura de la solución", level=1)
    add_fixed_table(
        doc,
        ["Capa", "Tecnologías y responsabilidad"],
        [
            ["Presentación", "React 19 + TypeScript; componentes, formularios, tablas, tarjetas y navegación responsiva."],
            ["Construcción", "Vinext/Vite; compilación de producción y servidor de desarrollo."],
            ["Integración", "Cliente HTTP configurable mediante NEXT_PUBLIC_API_URL; valor local por defecto: http://localhost:5001/api/v1."],
            ["Servicios", "API REST con Express, TypeScript y Mongoose; módulos de catálogo, caja, inventario, clientes, abastecimiento, acceso y reportes."],
            ["Datos", "MongoDB para persistencia; conjunto demo local únicamente como contingencia visual del frontend."],
        ],
        [2500, 6860],
        font_size=9.3,
    )


def add_quality_page(doc: Document):
    doc.add_page_break()
    add_kicker(doc, "Evidencia técnica")
    add_title(doc, "11  Integración y pruebas", size=24)
    add_body(
        doc,
        "El frontend intenta consultar la API real al iniciar y al pulsar “Actualizar información”. "
        "Si el servicio o su base de datos no responden, conserva la navegación y muestra una banda visible de "
        "“Modo demostración”. Por eso las capturas de este informe son reproducibles incluso sin MongoDB en ejecución."
    )
    doc.add_heading("Correspondencia entre módulos y servicios", level=2)
    add_fixed_table(
        doc,
        ["Módulo", "Servicios principales"],
        [
            ["Resumen", "Catálogo, cajas, ventas, clientes y reportes consolidados."],
            ["Punto de venta", "GET /checkout/registers; POST /checkout/transactions"],
            ["Catálogo", "GET/POST /catalog/articles; GET/POST /catalog/sections"],
            ["Inventario", "POST /inventory/movements; GET /reports/inventory"],
            ["Abastecimiento", "GET/POST /supply/suppliers; POST /supply/restock"],
            ["Clientes", "GET/POST/PUT /clients"],
            ["Equipo", "GET/POST /access/users; GET/POST /access/roles"],
            ["Reportes", "GET /reports/sales; /reports/inventory; /reports/clients"],
        ],
        [2500, 6860],
        font_size=8.8,
    )
    doc.add_heading("Resultados de verificación", level=2)
    add_fixed_table(
        doc,
        ["Comprobación", "Resultado"],
        [
            ["Backend", "15 suites y 46 pruebas aprobadas."],
            ["Frontend", "Compilación de producción completada; 2 pruebas de render e integración aprobadas."],
            ["Calidad estática", "ESLint ejecutado sin errores."],
            ["Recorrido visual", "8 de 8 módulos capturados y revisados en el navegador."],
        ],
        [3000, 6360],
    )
    add_callout(
        doc,
        "Lectura correcta del estado observado",
        "La banda amarilla de las capturas no representa un defecto del frontend: informa que la API local no estaba "
        "disponible durante la sesión de evidencia. Al levantar backend y MongoDB, la misma interfaz utiliza datos reales.",
        fill=AMBER_LIGHT,
        accent="8A5C00",
    )


def add_conclusions_page(doc: Document):
    doc.add_page_break()
    add_kicker(doc, "Cierre y defensa")
    add_title(doc, "12  Conclusiones y guion de exposición", size=24)
    add_body(
        doc,
        "El resultado cumple el objetivo central de la práctica: existe un frontend React funcional, organizado por "
        "procesos de negocio y preparado para consumir la API del minimarket. La solución conserva claridad operativa, "
        "comunica los estados de carga y error y mantiene la experiencia disponible mediante datos demostrativos."
    )
    add_fixed_table(
        doc,
        ["Criterio de la rúbrica", "Evidencia que debe mostrar el equipo"],
        [
            ["Repositorio documentado (2 pt)", "README, arquitectura, requisitos, tareas y este informe."],
            ["Trabajo colaborativo (2 pt)", "Historial con al menos 20 commits y participación de todos los integrantes."],
            ["Exposición ordenada (2 pt)", "Seguir el guion incluido abajo."],
            ["Aplicación funcional (4 pt)", "Demostrar navegación, formularios, una venta y reportes con la API activa."],
        ],
        [3300, 6060],
        font_size=9.2,
    )
    doc.add_heading("Guion sugerido para 5-7 minutos", level=2)
    add_fixed_table(
        doc,
        ["Tiempo", "Qué explicar o demostrar"],
        [
            ["0:00-0:45", "Problema: el minimarket necesita centralizar ventas, existencias y seguimiento comercial."],
            ["0:45-1:30", "Arquitectura: React se comunica con la API REST; MongoDB conserva la información."],
            ["1:30-3:30", "Demostración: Resumen -> Punto de venta -> Inventario -> Reportes."],
            ["3:30-4:30", "Módulos complementarios: Catálogo, Abastecimiento, Clientes y Equipo."],
            ["4:30-5:30", "Pruebas realizadas, comportamiento sin API y decisiones de alcance."],
            ["5:30-6:00", "Conclusión, trabajo futuro y preguntas."],
        ],
        [1900, 7460],
        font_size=9.0,
    )
    doc.add_heading("Datos finales por completar", level=2)
    add_fixed_table(
        doc,
        ["Campo", "Valor"],
        [
            ["Repositorio", "[PEGAR ENLACE DEFINITIVO]"],
            ["Grupo e integrantes", "[COMPLETAR]"],
            ["Asignatura y docente", "[COMPLETAR]"],
            ["Fecha de entrega", "[COMPLETAR]"],
        ],
        [3000, 6360],
    )


def build():
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "Frontend para la gestión de un minimarket"
    doc.core_properties.subject = "Informe técnico de práctica final"
    doc.core_properties.author = "Grupo XX"
    doc.core_properties.keywords = "React, minimarket, frontend, API REST, TypeScript"
    doc.core_properties.comments = "Informe editable generado para la entrega académica."

    add_cover(doc)
    add_contents_page(doc)
    add_summary_architecture_page(doc)

    modules = [
        dict(
            number="3",
            kicker="Centro de operaciones",
            title="Panel principal",
            lead="Resume la situación diaria del negocio antes de iniciar una tarea.",
            image="01-dashboard.png",
            caption="Panel con indicadores, ritmo de ventas y alertas de reposición.",
            purpose="Concentrar los indicadores más importantes y ofrecer accesos rápidos.",
            actions="Iniciar una nueva venta, abrir inventario, registrar un cliente o revisar reportes.",
            data="Ventas acumuladas, unidades en stock, clientes frecuentes, caja y productos con stock bajo.",
            figure_no=2,
        ),
        dict(
            number="4",
            kicker="Venta rápida",
            title="Punto de venta",
            lead="Permite construir una transacción desde el catálogo y enviarla al backend.",
            image="02-checkout.png",
            caption="Pantalla de punto de venta con catálogo y resumen de transacción.",
            purpose="Registrar una compra de forma directa desde la caja.",
            actions="Buscar productos, agregarlos, seleccionar caja y cliente, aplicar descuento y método de pago.",
            data="Productos disponibles, precio, stock, subtotal, descuento y total.",
            figure_no=3,
        ),
        dict(
            number="5",
            kicker="Maestro de productos",
            title="Catálogo",
            lead="Organiza la información comercial que alimenta ventas e inventario.",
            image="03-catalog.png",
            caption="Listado de productos y formulario de creación.",
            purpose="Consultar y registrar artículos con su sección, código, precio y existencia inicial.",
            actions="Buscar productos y crear nuevos registros desde un formulario lateral.",
            data="Nombre, sección, código de barras, precio y unidades disponibles.",
            figure_no=4,
        ),
        dict(
            number="6",
            kicker="Control de existencias",
            title="Inventario",
            lead="Hace visible el stock y registra sus entradas o salidas.",
            image="04-inventory.png",
            caption="Resumen de existencias, alertas y formulario de movimientos.",
            purpose="Evitar faltantes y mantener trazabilidad básica del inventario.",
            actions="Seleccionar un producto, elegir entrada o salida, indicar cantidad y registrar el movimiento.",
            data="Total de unidades, productos activos, stock bajo y disponibilidad por artículo.",
            figure_no=5,
        ),
        dict(
            number="7",
            kicker="Cadena de suministro",
            title="Abastecimiento",
            lead="Relaciona proveedores con las órdenes de reposición.",
            image="05-supply.png",
            caption="Directorio de proveedores y formulario de nueva reposición.",
            purpose="Registrar contactos comerciales y documentar compras para reponer inventario.",
            actions="Crear proveedor y generar una reposición seleccionando producto, cantidad y costo.",
            data="Empresa, contacto, teléfono, producto, cantidad y costo unitario.",
            figure_no=6,
        ),
        dict(
            number="8",
            kicker="Relación con clientes",
            title="Clientes",
            lead="Mantiene un directorio reutilizable durante futuras ventas.",
            image="06-clients.png",
            caption="Directorio de clientes y formulario de nuevo registro.",
            purpose="Conservar información básica para asociar compras y analizar recurrencia.",
            actions="Buscar, crear y preparar la edición de clientes.",
            data="Nombre, identificación, teléfono, correo y preferencias.",
            figure_no=7,
        ),
        dict(
            number="9",
            kicker="Organización del equipo",
            title="Usuarios y roles",
            lead="Representa la distribución de responsabilidades dentro del sistema.",
            image="07-access.png",
            caption="Usuarios activos, roles disponibles y formularios de alta.",
            purpose="Administrar perfiles de trabajo y roles operativos.",
            actions="Consultar usuarios/roles y crear registros de equipo.",
            data="Nombre, usuario, rol, descripción y estado.",
            figure_no=8,
            note="La gestión visual está implementada; la autenticación y la seguridad aplicada quedan fuera del alcance actual.",
        ),
        dict(
            number="10",
            kicker="Análisis del negocio",
            title="Reportes",
            lead="Presenta la información consolidada que entrega el backend.",
            image="08-reports.png",
            caption="Indicadores y detalle tabular de ventas.",
            purpose="Facilitar el análisis de ventas, inventario y clientes.",
            actions="Cambiar entre reportes y revisar el detalle de cada registro.",
            data="Venta acumulada, ticket promedio, valor del inventario, fechas, métodos de pago y totales.",
            figure_no=9,
        ),
    ]
    for module in modules:
        add_module_page(doc, **module)

    add_quality_page(doc)
    add_conclusions_page(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
